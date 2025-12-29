from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from datetime import datetime
from pathlib import Path


def add_header(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    return heading


def add_paragraph_with_style(doc, text, bold=False, italic=False):
    """添加段落并设置样式"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p


def create_test_case_table(doc, test_cases):
    """创建测试用例表格"""
    # 创建表格：Test ID, Description, Steps, Expected Result, Actual Result, Status, Comments
    table = doc.add_table(rows=1, cols=7)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    header_cells = table.rows[0].cells
    headers = [
        "Test ID",
        "Test Case",
        "Test Steps",
        "Expected Result",
        "Actual Result",
        "Status",
        "Comments",
    ]

    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = header
        # 设置表头样式
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
        # 设置背景色
        shading_elm = cell._element.get_or_add_tcPr()
        shading = parse_xml(r'<w:shd {} w:fill="4472C4"/>'.format(nsdecls("w")))
        shading_elm.append(shading)

    # 设置列宽
    widths = [
        Inches(0.6),
        Inches(1.5),
        Inches(2.0),
        Inches(1.5),
        Inches(1.2),
        Inches(0.8),
        Inches(1.4),
    ]
    for i, width in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = width

    # 添加测试用例
    for tc in test_cases:
        row_cells = table.add_row().cells
        row_cells[0].text = tc["id"]
        row_cells[1].text = tc["case"]
        row_cells[2].text = tc["steps"]
        row_cells[3].text = tc["expected"]
        row_cells[4].text = ""  # 留空给测试者填写
        row_cells[5].text = ""  # 留空给测试者填写
        row_cells[6].text = ""  # 留空给测试者填写

        # 设置单元格样式
        for cell in row_cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    return table


def parse_xml(xml_string):
    """解析 XML 字符串"""
    from docx.oxml import parse_xml as px

    return px(xml_string)


def nsdecls(*prefixes):
    """命名空间声明"""
    from docx.oxml.ns import nsdecls as nd

    return nd(*prefixes)


def generate_uat_document():
    """生成完整的 UAT 文档"""

    doc = Document()

    # ========================================
    # 文档属性
    # ========================================
    sections = doc.sections
    for section in sections:
        section.page_height = Inches(11.69)  # A4
        section.page_width = Inches(8.27)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # ========================================
    # 封面
    # ========================================
    title = doc.add_heading("User Acceptance Testing (UAT)", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("WinScope - Digital Forensics Evidence Collection Tool")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(68, 114, 196)

    version = doc.add_paragraph("Version 0.2.0")
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_run = version.runs[0]
    version_run.font.size = Pt(14)

    doc.add_paragraph()  # 空行
    doc.add_paragraph()

    # 文档信息表
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = "Light Grid Accent 1"

    info_data = [
        ("Document Title:", "User Acceptance Testing Plan"),
        ("Project Name:", "WinScope - Digital Forensics Tool"),
        ("Version:", "0.2.0"),
        ("Date:", datetime.now().strftime("%Y-%m-%d")),
        ("Prepared By:", "Alfred (Student)"),
    ]

    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_page_break()

    # ========================================
    # 1. Document Control
    # ========================================
    add_header(doc, "1. Document Control", level=1)

    doc.add_paragraph(
        "This document outlines the User Acceptance Testing (UAT) procedures for WinScope v0.2.0, "
        "a digital forensics evidence collection tool developed as part of a Final Year Project (FYP)."
    )

    # Version History
    add_header(doc, "1.1 Version History", level=2)

    version_table = doc.add_table(rows=2, cols=4)
    version_table.style = "Light Grid Accent 1"

    # 表头
    headers = ["Version", "Date", "Author", "Changes"]
    for i, header in enumerate(headers):
        cell = version_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    # 数据
    version_table.rows[1].cells[0].text = "0.2.0"
    version_table.rows[1].cells[1].text = datetime.now().strftime("%Y-%m-%d")
    version_table.rows[1].cells[2].text = "Alfred"
    version_table.rows[1].cells[3].text = "Initial UAT document"

    doc.add_paragraph()

    # ========================================
    # 2. Introduction
    # ========================================
    add_header(doc, "2. Introduction", level=1)

    add_header(doc, "2.1 Purpose", level=2)
    doc.add_paragraph(
        "The purpose of this User Acceptance Testing is to verify that WinScope meets the "
        "functional and non-functional requirements for digital forensics evidence collection "
        "in Windows incident response scenarios."
    )

    add_header(doc, "2.2 Scope", level=2)
    doc.add_paragraph("This UAT covers the following areas:")

    scope_items = [
        "User Interface (UI) functionality and usability",
        "Evidence collection modules (Memory, Disk, Network, Registry, etc.)",
        "Configuration management and settings",
        "Report generation and output formats",
        "Error handling and system stability",
        "Compatibility with forensic analysis tools (Volatility, Autopsy)",
    ]

    for item in scope_items:
        doc.add_paragraph(item, style="List Bullet")

    add_header(doc, "2.3 Test Environment", level=2)

    env_table = doc.add_table(rows=6, cols=2)
    env_table.style = "Light Grid Accent 1"

    env_data = [
        ("Operating System:", "Windows 10/11 (64-bit)"),
        ("RAM:", "Minimum 4GB"),
        ("Disk Space:", "Minimum 10GB free"),
        ("Privileges:", "Administrator rights"),
        ("Additional Software:", "WinPmem (optional for memory acquisition)"),
        ("Test Duration:", "2-3 hours"),
    ]

    for i, (label, value) in enumerate(env_data):
        row = env_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_page_break()

    # ========================================
    # 3. Test Cases
    # ========================================
    add_header(doc, "3. Test Cases", level=1)

    doc.add_paragraph(
        "Please execute each test case and record the actual results. "
        "Mark the status as PASS, FAIL, or BLOCKED, and provide comments if necessary."
    )

    doc.add_paragraph()

    # 3.1 Application Launch & UI
    add_header(doc, "3.1 Application Launch & User Interface", level=2)

    ui_test_cases = [
        {
            "id": "TC-01",
            "case": "Application Launch",
            "steps": '1. Right-click WinScope.exe\n2. Select "Run as Administrator"\n3. Observe application startup',
            "expected": "Application launches successfully\nMain window appears within 5 seconds\nNo error messages displayed",
        },
        {
            "id": "TC-02",
            "case": "UI Layout",
            "steps": "1. Observe main window layout\n2. Check all panels are visible\n3. Verify responsiveness",
            "expected": "All panels (Modules, Configuration, Progress & Logs) are visible\nUI is responsive and readable\nNo visual glitches",
        },
        {
            "id": "TC-03",
            "case": "Module Selection",
            "steps": '1. Click "Select All" button\n2. Click "Deselect All" button\n3. Manually select individual modules',
            "expected": "All modules selected/deselected correctly\nCheckboxes respond immediately\nModule count updates accurately",
        },
        {
            "id": "TC-04",
            "case": "System Information Display",
            "steps": "1. Check System Information card\n2. Verify displayed information",
            "expected": "Hostname, IP, OS version, username, and timestamp displayed correctly\nAll information is accurate",
        },
        {
            "id": "TC-05",
            "case": "Privilege Check",
            "steps": "1. Observe privilege status indicator\n2. Check module availability",
            "expected": "Administrator status shown clearly\nModules requiring admin are enabled\nPrivilege indicator is accurate",
        },
    ]

    create_test_case_table(doc, ui_test_cases)

    doc.add_page_break()

    # 3.2 Configuration
    add_header(doc, "3.2 Configuration Management", level=2)

    config_test_cases = [
        {
            "id": "TC-06",
            "case": "Output Directory Selection",
            "steps": '1. Click "Browse" button\n2. Select a directory\n3. Verify path is updated',
            "expected": "File browser opens correctly\nSelected path appears in text field\nPath is valid and accessible",
        },
        {
            "id": "TC-07",
            "case": "Compression Options",
            "steps": '1. Select "None" compression\n2. Select "ZIP" compression\n3. Select "7-Zip" compression',
            "expected": "All compression options are selectable\nOnly one option selected at a time\nSelection persists",
        },
        {
            "id": "TC-08",
            "case": "Hash Algorithm Selection",
            "steps": "1. Check MD5\n2. Check SHA-1\n3. Check SHA-256\n4. Uncheck all",
            "expected": "All hash algorithms can be selected\nMultiple selections allowed\nAt least one must be selected for collection",
        },
        {
            "id": "TC-09",
            "case": "Save Configuration Profile",
            "steps": '1. Configure settings\n2. Click "Save Profile"\n3. Enter profile name\n4. Save',
            "expected": "Profile saved successfully\nConfirmation message displayed\nProfile can be retrieved later",
        },
        {
            "id": "TC-10",
            "case": "Load Configuration Profile",
            "steps": '1. Click "Load Profile"\n2. Select saved profile\n3. Verify settings loaded',
            "expected": "Profile list displays correctly\nAll settings restored accurately\nUI updates to reflect loaded settings",
        },
    ]

    create_test_case_table(doc, config_test_cases)

    doc.add_page_break()

    # 3.3 Evidence Collection
    add_header(doc, "3.3 Evidence Collection Modules", level=2)

    collection_test_cases = [
        {
            "id": "TC-11",
            "case": "Memory Acquisition",
            "steps": "1. Select Memory module\n2. Start collection\n3. Monitor progress\n4. Check output",
            "expected": "Memory dump file created (.raw)\nProgress bar updates in real-time\nHash values calculated\nMetadata files generated",
        },
        {
            "id": "TC-12",
            "case": "Disk Acquisition",
            "steps": "1. Select Disk module\n2. Start collection\n3. Monitor progress\n4. Check output",
            "expected": "File list CSV created for each partition\nProgress updates per partition\nDisk information JSON generated\nAutopsy guide created",
        },
        {
            "id": "TC-13",
            "case": "Network Collection",
            "steps": "1. Select Network module\n2. Start collection\n3. Check output files",
            "expected": "Network connections captured\nARP cache exported\nRouting tables saved\nDNS cache collected",
        },
        {
            "id": "TC-14",
            "case": "Registry Export",
            "steps": "1. Select Registry module\n2. Start collection\n3. Verify registry hives",
            "expected": "All major registry hives exported\nSAM, SYSTEM, SOFTWARE, SECURITY\nFiles are valid registry format",
        },
        {
            "id": "TC-15",
            "case": "Event Logs Collection",
            "steps": "1. Select Event Logs module\n2. Start collection\n3. Check log files",
            "expected": "System, Application, Security logs copied\nLog files in .evtx format\nLog metadata recorded",
        },
        {
            "id": "TC-16",
            "case": "Browser Artifacts",
            "steps": "1. Select Browser module\n2. Start collection\n3. Review artifacts",
            "expected": "Browser history collected\nCookies and cache captured\nBookmarks exported\nMultiple browsers supported",
        },
        {
            "id": "TC-17",
            "case": "Live System Info",
            "steps": "1. Select Live System module\n2. Start collection\n3. Check system data",
            "expected": "Running processes listed\nStartup items recorded\nScheduled tasks exported\nInstalled programs cataloged",
        },
        {
            "id": "TC-18",
            "case": "File System Metadata",
            "steps": "1. Select File System module\n2. Start collection\n3. Verify metadata",
            "expected": "MFT records extracted\nFile metadata (MAC times) recorded\nAlternate data streams detected",
        },
    ]

    create_test_case_table(doc, collection_test_cases)

    doc.add_page_break()

    # 3.4 Progress & Logging
    add_header(doc, "3.4 Progress Monitoring & Logging", level=2)

    progress_test_cases = [
        {
            "id": "TC-19",
            "case": "Overall Progress Bar",
            "steps": "1. Start collection with multiple modules\n2. Observe overall progress bar",
            "expected": "Progress bar shows X/Y modules completed\nPercentage updates accurately\nBar fills smoothly",
        },
        {
            "id": "TC-20",
            "case": "Current Module Progress",
            "steps": "1. Monitor current module progress\n2. Check percentage updates",
            "expected": "Current module name displayed\nProgress bar shows 0-100%\nUpdates in real-time",
        },
        {
            "id": "TC-21",
            "case": "Log Viewer",
            "steps": "1. Monitor log viewer during collection\n2. Check log messages",
            "expected": "Log messages appear in real-time\nColor coding (INFO, WARNING, ERROR, SUCCESS)\nMessages are clear and informative",
        },
        {
            "id": "TC-22",
            "case": "Export Logs",
            "steps": '1. Click "Export Logs" button\n2. Save log file\n3. Open exported file',
            "expected": "Log export dialog appears\nFile saved successfully\nExported log is readable",
        },
        {
            "id": "TC-23",
            "case": "View Detailed Log",
            "steps": '1. Complete collection\n2. Click "View Detailed Log"\n3. Check log file',
            "expected": "Detailed log opens in Notepad\nAll collection activities logged\nTimestamps are accurate",
        },
    ]

    create_test_case_table(doc, progress_test_cases)

    doc.add_page_break()

    # 3.5 Report Generation
    add_header(doc, "3.5 Report Generation & Output", level=2)

    report_test_cases = [
        {
            "id": "TC-24",
            "case": "HTML Report Generation",
            "steps": '1. Complete collection\n2. Click "View Report"\n3. Review HTML report',
            "expected": "Report opens in web browser\nAll sections present (Summary, Modules, Files)\nProfessional formatting",
        },
        {
            "id": "TC-25",
            "case": "Report Content Accuracy",
            "steps": "1. Review report contents\n2. Verify against actual collection",
            "expected": "All collected files listed\nHash values correct\nTimestamps accurate\nModule statuses correct",
        },
        {
            "id": "TC-26",
            "case": "Evidence Package Structure",
            "steps": "1. Navigate to output directory\n2. Review folder structure\n3. Check file organization",
            "expected": "Clean folder structure\nModules in separate folders\nLogs in root\nReport in root",
        },
        {
            "id": "TC-27",
            "case": "Compression",
            "steps": "1. Enable ZIP compression\n2. Complete collection\n3. Check compressed file",
            "expected": "ZIP file created\nAll evidence included\nFile size reasonable\nArchive is valid",
        },
        {
            "id": "TC-28",
            "case": "Hash Verification",
            "steps": "1. Enable hash verification\n2. Complete collection\n3. Check hash files",
            "expected": "Hash files generated\nMD5, SHA-1, SHA-256 values present\nHashes can be verified manually",
        },
    ]

    create_test_case_table(doc, report_test_cases)

    doc.add_page_break()

    # 3.6 Error Handling
    add_header(doc, "3.6 Error Handling & Edge Cases", level=2)

    error_test_cases = [
        {
            "id": "TC-29",
            "case": "Insufficient Disk Space",
            "steps": "1. Select output to nearly full drive\n2. Start collection\n3. Observe behavior",
            "expected": "Warning message displayed\nCollection continues if possible\nGraceful failure with error message",
        },
        {
            "id": "TC-30",
            "case": "Missing WinPmem",
            "steps": "1. Remove WinPmem from tools folder\n2. Select Memory module\n3. Start collection",
            "expected": "Warning message about missing tool\nInstructions provided\nOther modules continue normally",
        },
        {
            "id": "TC-31",
            "case": "Stop Collection",
            "steps": '1. Start collection\n2. Click "Stop Collection"\n3. Confirm stop',
            "expected": "Confirmation dialog appears\nCollection stops gracefully\nPartial results saved",
        },
        {
            "id": "TC-32",
            "case": "Invalid Output Path",
            "steps": "1. Enter invalid path (e.g., Z:\\nonexistent)\n2. Start collection",
            "expected": "Error message displayed\nUser prompted to select valid path\nCollection does not start",
        },
        {
            "id": "TC-33",
            "case": "No Modules Selected",
            "steps": '1. Deselect all modules\n2. Click "Start Collection"',
            "expected": 'Warning message displayed\n"Please select at least one module"\nCollection does not start',
        },
    ]

    create_test_case_table(doc, error_test_cases)

    doc.add_page_break()

    # 3.7 Usability & Performance
    add_header(doc, "3.7 Usability & Performance", level=2)

    usability_test_cases = [
        {
            "id": "TC-34",
            "case": "UI Responsiveness",
            "steps": "1. Interact with UI during collection\n2. Try clicking buttons\n3. Test window resizing",
            "expected": "UI remains responsive\nButtons clickable (even if disabled)\nWindow resizes smoothly\nNo freezing",
        },
        {
            "id": "TC-35",
            "case": "Collection Performance",
            "steps": "1. Time a full collection\n2. Monitor CPU and memory usage",
            "expected": "Collection completes in reasonable time\nCPU usage acceptable (<80%)\nMemory usage stable\nNo crashes",
        },
        {
            "id": "TC-36",
            "case": "Multiple Collections",
            "steps": "1. Run collection\n2. Complete\n3. Run another collection immediately",
            "expected": "Second collection starts successfully\nNo errors from previous collection\nProgress resets correctly",
        },
        {
            "id": "TC-37",
            "case": "Help & Documentation",
            "steps": '1. Click "Help" menu\n2. View documentation\n3. Check tooltips',
            "expected": "Help content is clear and useful\nDocumentation accessible\nTooltips provide guidance",
        },
        {
            "id": "TC-38",
            "case": "Application Exit",
            "steps": "1. Close application during idle\n2. Close during collection\n3. Check cleanup",
            "expected": "Clean exit during idle\nWarning during collection\nNo temp files left behind",
        },
    ]

    create_test_case_table(doc, usability_test_cases)

    doc.add_page_break()

    # ========================================
    # 4. Forensic Tool Compatibility
    # ========================================
    add_header(doc, "4. Forensic Tool Compatibility Testing", level=1)

    doc.add_paragraph(
        "These tests verify that collected evidence can be imported and analyzed in "
        "industry-standard forensic tools."
    )

    doc.add_paragraph()

    add_header(doc, "4.1 Volatility Compatibility", level=2)

    volatility_test_cases = [
        {
            "id": "TC-39",
            "case": "Memory Dump Format",
            "steps": "1. Collect memory dump\n2. Open in Volatility\n3. Run imageinfo command",
            "expected": "Memory dump loads successfully\nProfile detected correctly\nNo format errors",
        },
        {
            "id": "TC-40",
            "case": "Volatility Analysis",
            "steps": "1. Run basic Volatility plugins\n2. Check process list\n3. Verify output",
            "expected": "Plugins execute successfully\nProcess information extracted\nResults are valid",
        },
    ]

    create_test_case_table(doc, volatility_test_cases)

    doc.add_paragraph()

    add_header(doc, "4.2 Autopsy Compatibility", level=2)

    autopsy_test_cases = [
        {
            "id": "TC-41",
            "case": "Disk Data Import",
            "steps": "1. Open Autopsy\n2. Create new case\n3. Import collected file lists",
            "expected": "File lists import successfully\nMetadata preserved\nTimestamps accurate",
        },
        {
            "id": "TC-42",
            "case": "File System Analysis",
            "steps": "1. Analyze imported data in Autopsy\n2. View file system timeline\n3. Check artifacts",
            "expected": "Timeline generated correctly\nFile artifacts detected\nMAC times preserved",
        },
    ]

    create_test_case_table(doc, autopsy_test_cases)

    doc.add_page_break()

    # ========================================
    # 5. Test Results Summary
    # ========================================
    add_header(doc, "5. Test Results Summary", level=1)

    doc.add_paragraph("Please complete this section after finishing all test cases.")

    doc.add_paragraph()

    # Summary Table
    summary_table = doc.add_table(rows=7, cols=2)
    summary_table.style = "Light Grid Accent 1"

    summary_labels = [
        "Total Test Cases:",
        "Tests Passed:",
        "Tests Failed:",
        "Tests Blocked:",
        "Pass Rate:",
        "Testing Date:",
        "Tester Name:",
    ]

    for i, label in enumerate(summary_labels):
        row = summary_table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[1].text = "_________________"

    doc.add_paragraph()

    # ========================================
    # 6. Issues & Defects Log
    # ========================================
    add_header(doc, "6. Issues & Defects Log", level=1)

    doc.add_paragraph("Please document any issues, bugs, or defects encountered during testing.")

    doc.add_paragraph()

    # Issues Table
    issues_table = doc.add_table(rows=6, cols=5)
    issues_table.style = "Light Grid Accent 1"

    # 表头
    issue_headers = ["Issue ID", "Test Case", "Description", "Severity", "Status"]
    for i, header in enumerate(issue_headers):
        cell = issues_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    # 示例行
    example_row = issues_table.rows[1]
    example_row.cells[0].text = "ISS-01"
    example_row.cells[1].text = "TC-XX"
    example_row.cells[2].text = "Description of issue..."
    example_row.cells[3].text = "High/Medium/Low"
    example_row.cells[4].text = "Open/Resolved"

    # 空行供填写
    for i in range(2, 6):
        row = issues_table.rows[i]
        for cell in row.cells:
            cell.text = ""

    doc.add_paragraph()

    # ========================================
    # 7. Overall Assessment
    # ========================================
    add_header(doc, "7. Overall Assessment", level=1)

    add_header(doc, "7.1 Strengths", level=2)
    doc.add_paragraph("List the application's strengths:")
    for i in range(5):
        doc.add_paragraph(f"{i+1}. _________________________________________________")

    doc.add_paragraph()

    add_header(doc, "7.2 Areas for Improvement", level=2)
    doc.add_paragraph("List areas that need improvement:")
    for i in range(5):
        doc.add_paragraph(f"{i+1}. _________________________________________________")

    doc.add_paragraph()

    add_header(doc, "7.3 Overall Recommendation", level=2)

    recommendation_table = doc.add_table(rows=4, cols=2)
    recommendation_table.style = "Light Grid Accent 1"

    recommendations = [
        ("☐ Approved for Production", "Application meets all requirements and is ready for use"),
        ("☐ Approved with Minor Issues", "Application is acceptable with documented minor issues"),
        ("☐ Not Approved - Major Issues", "Significant issues require resolution before approval"),
        ("☐ Requires Re-testing", "Issues found require fixes and additional testing"),
    ]

    for i, (option, description) in enumerate(recommendations):
        row = recommendation_table.rows[i]
        row.cells[0].text = option
        row.cells[1].text = description
        row.cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()

    # ========================================
    # 8. Sign-off
    # ========================================
    add_header(doc, "8. Sign-off", level=1)

    signoff_table = doc.add_table(rows=4, cols=2)
    signoff_table.style = "Light Grid Accent 1"

    signoff_data = [
        ("Tester Name:", "_______________________________________"),
        ("Signature:", "_______________________________________"),
        ("Date:", "_______________________________________"),
        ("Comments:", "_______________________________________"),
    ]

    for i, (label, value) in enumerate(signoff_data):
        row = signoff_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer note
    footer_para = doc.add_paragraph(
        "Thank you for participating in the User Acceptance Testing of WinScope. "
        "Your feedback is invaluable for improving the quality and usability of this tool."
    )
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.runs[0].font.italic = True
    footer_para.runs[0].font.size = Pt(10)
    footer_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # ========================================
    # 保存文档
    # ========================================
    output_path = Path("WinScope_UAT_Document_v0.2.0.docx")
    doc.save(output_path)

    print(f"✓ UAT Document generated successfully: {output_path}")
    return output_path


if __name__ == "__main__":
    print("=" * 60)
    print(" WinScope UAT Document Generator")
    print("=" * 60)
    print()

    try:
        output_file = generate_uat_document()
        print()
        print("=" * 60)
        print(" Success!")
        print("=" * 60)
        print(f" Document: {output_file}")
        print(f" Size: {output_file.stat().st_size / 1024:.2f} KB")
        print()
        print(" You can now send this document to your teacher for UAT.")
        print("=" * 60)
    except Exception as e:
        print(f"✗ Error generating document: {e}")
        import traceback

        traceback.print_exc()
